"""
Handles .docx files via python-docx.
Extracts paragraphs and tables, preserving structure.
"""

import os
import uuid
import time
from core.document import Document, DocumentPage, Table, make_document
from core.config import Config
from core.logger import get_logger
from core.errors import ParseError
from parsers.base import BaseParser
from docx import Document as DocxDocument

logger = get_logger("docx")


class DocxParser(BaseParser):

    def get_name(self) -> str:
        return "docx"

    def can_handle(self, file_path: str) -> bool:
        return os.path.splitext(file_path)[1].lower() == ".docx"

    def is_available(self, config: Config) -> bool:
        return True

    def parse(self, file_path: str, config: Config) -> Document:
        start = time.time()
        file_name = os.path.basename(file_path)

        try:
            doc = DocxDocument(file_path)
        except Exception as e:
            raise ParseError(
                f"Could not open DOCX: {e}",
                file_name=file_name,
                retryable=False,
            )

        # Extract paragraphs
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())

        # Extract tables as structured Table objects
        tables = []
        for tbl in doc.tables:
            headers = []
            rows = []
            for row_idx, row in enumerate(tbl.rows):
                cells = [cell.text.strip() for cell in row.cells]
                cells = [c for c in cells if c]  # drop empty cells
                if not cells:
                    continue
                if row_idx == 0:
                    headers = cells
                else:
                    rows.append(cells)

                # Also include table text in body
                text_parts.append(" | ".join(cells))

            if headers:
                tables.append(Table(
                    page_num=1,
                    title="",
                    headers=headers,
                    rows=rows,
                    cells=[],
                    raw_text="\n".join(" | ".join(r) for r in [headers] + rows),
                ))

        full_text = "\n\n".join(text_parts)
        if not full_text.strip():
            raise ParseError(
                "DOCX contains no extractable text",
                file_name=file_name,
                retryable=False,
            )

        # DOCX has no native pages — chunk into ~3000-char logical pages
        pages = []
        chunk_size = 3000
        chunks = [full_text[i:i + chunk_size].strip()
                  for i in range(0, len(full_text), chunk_size)
                  if full_text[i:i + chunk_size].strip()]

        for page_num, chunk in enumerate(chunks, start=1):
            # Attach tables to first page only
            page_tables = tables if page_num == 1 else []
            pages.append(DocumentPage(
                page_num=page_num,
                text=chunk,
                tables=page_tables,
                images=[],
                layout=[],
                entities=[],
                word_count=len(chunk.split()),
                ocr_confidence=1.0,
            ))

        duration_ms = int((time.time() - start) * 1000)
        logger.info("DOCX parse complete",
                    file=file_name,
                    pages=len(pages),
                    tables=len(tables),
                    duration_ms=duration_ms)

        return make_document(
            id=str(uuid.uuid4()),
            file_name=file_name,
            file_type=".docx",
            file_path=file_path,
            pages=pages,
            parser_used="docx",
            metadata={
                "is_scanned": False,
                "parse_duration_ms": duration_ms,
                "file_size_bytes": os.path.getsize(file_path),
            },
        )