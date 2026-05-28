import os
import re
import threading
from dotenv import load_dotenv
from llama_index.core.node_parser import SentenceSplitter
from llama_index.core import Document
from db import insert_document, insert_chunks

load_dotenv()

embed_model = None
_model_lock = threading.Lock()

def get_embed_model():
    global embed_model
    if embed_model is None:
        with _model_lock:
            if embed_model is None:
                print("Loading embedding model... (first time only)")
                from llama_index.embeddings.huggingface import HuggingFaceEmbedding
                embed_model = HuggingFaceEmbedding(
                    model_name="BAAI/bge-small-en-v1.5",
                    device="cpu"
                )
    return embed_model

splitter = SentenceSplitter(chunk_size=512, chunk_overlap=64)

SUPPORTED_EXTENSIONS = [".pdf", ".docx", ".txt", ".csv", ".xlsx", ".rtf", ".md"]

# --- Parsers ---

def parse_pdf(file_path: str, use_llamaparse: bool = True) -> list[dict]:
    if use_llamaparse and os.getenv("LLAMA_CLOUD_API_KEY"):
        try:
            from llama_parse import LlamaParse
            parser = LlamaParse(
                api_key=os.getenv("LLAMA_CLOUD_API_KEY"),
                result_type="markdown",
                verbose=False
            )
            docs = parser.load_data(file_path)
            pages = [{"text": d.text.strip(), "page": str(i+1)} for i, d in enumerate(docs) if d.text.strip()]
            if pages:
                return pages
        except Exception as e:
            print(f"LlamaParse failed: {e} — falling back to pypdf")

    from pypdf import PdfReader
    reader = PdfReader(file_path)
    pages = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text and text.strip():
            pages.append({"text": text.strip(), "page": str(i+1)})
    return pages


def parse_docx(file_path: str) -> list[dict]:
    from docx import Document as DocxDocument
    doc = DocxDocument(file_path)
    current_text = []
    page_num = 1

    for para in doc.paragraphs:
        if para.text.strip():
            current_text.append(para.text.strip())

    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            current_text.append("\n".join(rows))

    full_text = "\n\n".join(current_text)
    chunks = []
    chunk_size = 3000
    for i in range(0, len(full_text), chunk_size):
        chunk = full_text[i:i+chunk_size].strip()
        if chunk:
            chunks.append({"text": chunk, "page": str(page_num)})
            page_num += 1
    return chunks


def parse_txt(file_path: str) -> list[dict]:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()
    if not text.strip():
        return []
    chunks = []
    chunk_size = 3000
    for i in range(0, len(text), chunk_size):
        chunk = text[i:i+chunk_size].strip()
        if chunk:
            chunks.append({"text": chunk, "page": str(i // chunk_size + 1)})
    return chunks


def parse_md(file_path: str) -> list[dict]:
    return parse_txt(file_path)


def parse_csv(file_path: str) -> list[dict]:
    import pandas as pd
    df = pd.read_csv(file_path)
    chunks = []
    batch_size = 50
    for i in range(0, len(df), batch_size):
        batch = df.iloc[i:i+batch_size]
        text = f"Columns: {', '.join(df.columns.tolist())}\n\n"
        text += batch.to_string(index=False)
        chunks.append({
            "text": text,
            "page": f"rows {i+1}-{min(i+batch_size, len(df))}"
        })
    return chunks


def parse_xlsx(file_path: str) -> list[dict]:
    import pandas as pd
    chunks = []
    xl = pd.ExcelFile(file_path)
    for sheet_name in xl.sheet_names:
        df = pd.read_excel(file_path, sheet_name=sheet_name)
        if df.empty:
            continue
        batch_size = 50
        for i in range(0, len(df), batch_size):
            batch = df.iloc[i:i+batch_size]
            text = f"Sheet: {sheet_name}\nColumns: {', '.join(df.columns.astype(str).tolist())}\n\n"
            text += batch.to_string(index=False)
            chunks.append({
                "text": text,
                "page": f"{sheet_name} rows {i+1}-{min(i+batch_size, len(df))}"
            })
    return chunks


def parse_rtf(file_path: str) -> list[dict]:
    from striprtf.striprtf import rtf_to_text
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        rtf_content = f.read()
    text = rtf_to_text(rtf_content)
    return parse_txt_from_string(text)


def parse_txt_from_string(text: str) -> list[dict]:
    chunks = []
    chunk_size = 3000
    for i in range(0, len(text), chunk_size):
        chunk = text[i:i+chunk_size].strip()
        if chunk:
            chunks.append({"text": chunk, "page": str(i // chunk_size + 1)})
    return chunks


# --- File router ---

def parse_file(file_path: str, use_llamaparse: bool = True) -> list[dict]:
    ext = os.path.splitext(file_path)[1].lower()
    parsers = {
        ".pdf":  lambda: parse_pdf(file_path, use_llamaparse),
        ".docx": lambda: parse_docx(file_path),
        ".txt":  lambda: parse_txt(file_path),
        ".md":   lambda: parse_md(file_path),
        ".csv":  lambda: parse_csv(file_path),
        ".xlsx": lambda: parse_xlsx(file_path),
        ".rtf":  lambda: parse_rtf(file_path),
    }
    parser = parsers.get(ext)
    if not parser:
        return []
    return parser()


# --- Main file ingestion ---

def ingest_file(file_path: str, use_llamaparse: bool = True) -> dict:
    file_name = os.path.basename(file_path)
    ext = os.path.splitext(file_path)[1].lower()
    model = get_embed_model()

    if ext not in SUPPORTED_EXTENSIONS:
        return {"error": f"Unsupported file type: {ext}. Supported: {', '.join(SUPPORTED_EXTENSIONS)}"}

    print(f"Parsing {file_name} as {ext}")
    pages = parse_file(file_path, use_llamaparse)

    if not pages:
        return {"error": "Could not extract text. File may be empty or image-only."}

    doc_id = insert_document(file_name)
    chunk_rows = []

    for page_data in pages:
        doc = Document(text=page_data["text"])
        nodes = splitter.get_nodes_from_documents([doc])
        for node in nodes:
            clean_text = node.text.replace("\x00", " ").strip()
            if not clean_text:
                continue
            embedding = model.get_text_embedding(clean_text)
            chunk_rows.append({
                "document_id": doc_id,
                "content": clean_text,
                "embedding": embedding,
                "metadata": {
                    "page": page_data["page"],
                    "file": file_name
                }
            })

    insert_chunks(chunk_rows)

    return {
        "document_id": doc_id,
        "file": file_name,
        "chunks_stored": len(chunk_rows),
        "parser": ext.replace(".", "")
    }


# --- URL ingestion ---

def fetch_url_content(url: str) -> tuple[str, str]:
    """Returns (title, text) — uses Wikipedia API for Wikipedia URLs"""
    import httpx
    from bs4 import BeautifulSoup

    # Wikipedia handler
    wiki_match = re.search(r'wikipedia\.org/wiki/(.+)', url)
    if wiki_match:
        page_title = wiki_match.group(1).split("#")[0]  # remove anchor
        try:
            import httpx
            r = httpx.get(
                "https://en.wikipedia.org/w/api.php",
                params={
                    "action": "query",
                    "titles": page_title,
                    "prop": "extracts",
                    "explaintext": "1",
                    "exsectionformat": "plain",
                    "format": "json",
                    "redirects": "1"
                },
                headers={
                    "User-Agent": "DocIntel/1.0 (document intelligence tool; contact@docintel.app)"
                },
                timeout=20
            )
            if r.status_code == 200 and r.text.strip():
                data = r.json()
                pages = data.get("query", {}).get("pages", {})
                page = next(iter(pages.values()))
                if page.get("pageid"):  # valid page
                    title = page.get("title", page_title)
                    text = page.get("extract", "")
                    if text:
                        return title, text
        except Exception as e:
            print(f"Wikipedia API failed: {e}")
    # Fall through to generic handler if Wikipedia API fails

    # Generic handler
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
        "Accept-Language": "en-US,en;q=0.5",
        "Connection": "keep-alive",
    }
    try:
        response = httpx.get(url, headers=headers, timeout=15, follow_redirects=True)
        response.raise_for_status()
    except Exception as e:
        return "", f"Could not fetch URL: {str(e)}"

    soup = BeautifulSoup(response.text, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header", "aside", "form"]):
        tag.decompose()

    title = soup.title.string.strip() if soup.title else url
    main = soup.find("article") or soup.find("main") or soup.find("body")
    text = main.get_text(separator="\n") if main else soup.get_text(separator="\n")
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    return title, "\n".join(lines)


def ingest_url(url: str) -> dict:
    title, text = fetch_url_content(url)

    if text.startswith("Could not fetch") or text.startswith("Wikipedia fetch"):
        return {"error": text}

    if not text.strip():
        return {"error": "Could not extract text from this URL."}

    safe_title = f"{title[:60].replace('/', '-')}.url"
    model = get_embed_model()
    doc_id = insert_document(safe_title)

    chunk_size = 3000
    chunk_rows = []
    for i in range(0, len(text), chunk_size):
        chunk = text[i:i + chunk_size].strip()
        if not chunk:
            continue
        doc = Document(text=chunk)
        nodes = splitter.get_nodes_from_documents([doc])
        for node in nodes:
            clean = node.text.replace("\x00", " ").strip()
            if not clean:
                continue
            embedding = model.get_text_embedding(clean)
            chunk_rows.append({
                "document_id": doc_id,
                "content": clean,
                "embedding": embedding,
                "metadata": {
                    "page": str(i // chunk_size + 1),
                    "file": safe_title,
                    "source_url": url
                }
            })

    insert_chunks(chunk_rows)

    return {
        "document_id": doc_id,
        "file": safe_title,
        "title": title,
        "chunks_stored": len(chunk_rows),
        "parser": "url"
    }